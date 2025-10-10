"""
Text extraction module for different file formats
"""
import PyPDF2
import docx
from pathlib import Path
import logging
from striprtf.striprtf import rtf_to_text
from odf import text as odf_text
from odf.opendocument import load as odf_load
import markdown

logger = logging.getLogger(__name__)

class TextExtractor:
    """Handles text extraction from various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.rtf', '.odt', '.md']
    
    def extract_text(self, file_path):
        """
        Extract text from a file based on its extension
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            dict: {'success': bool, 'text': str, 'error': str}
        """
        try:
            file_path = Path(file_path)
            extension = file_path.suffix.lower()
            
            if extension == '.pdf':
                return self._extract_pdf(file_path)
            elif extension == '.docx':
                return self._extract_docx(file_path)
            elif extension == '.txt':
                return self._extract_txt(file_path)
            elif extension == '.rtf':
                return self._extract_rtf(file_path)
            elif extension == '.odt':
                return self._extract_odt(file_path)
            elif extension == '.md':
                return self._extract_markdown(file_path)
            else:
                return {
                    'success': False,
                    'text': '',
                    'error': f'Unsupported file format: {extension}'
                }
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return {
                'success': False,
                'text': '',
                'error': str(e)
            }
    
    def _extract_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            if not text.strip():
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text could be extracted from PDF'
                }
            
            return {
                'success': True,
                'text': text.strip(),
                'error': ''
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'PDF extraction error: {str(e)}'
            }
    
    def _extract_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                return {
                    'success': False,
                    'text': '',
                    'error': 'No text could be extracted from DOCX'
                }
            
            return {
                'success': True,
                'text': text.strip(),
                'error': ''
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'DOCX extraction error: {str(e)}'
            }
    
    def _extract_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    
                    if text.strip():
                        return {
                            'success': True,
                            'text': text.strip(),
                            'error': ''
                        }
                except UnicodeDecodeError:
                    continue
            
            return {
                'success': False,
                'text': '',
                'error': 'Could not decode text file with any supported encoding'
            }
            
        except Exception as e:
            return {
                'success': False,
                'text': '',
                'error': f'TXT extraction error: {str(e)}'
            }
